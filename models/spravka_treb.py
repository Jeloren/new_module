import base64
import io
from odoo import _, fields, models, api
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

class SpravkaTreb(models.Model):
    _name = "new_module.spravka_treb"
    _description = "This is my spravka_treb"
    
    f_student = fields.Many2one(comodel_name="new_module.students", string="Студент")
    display_name = fields.Char(compute = "Method")
    group_creating_date = fields.Date(string="Дата создания группы", compute="_compute_stud")
    stud_birth_day = fields.Date(string="День рождения", compute="_compute_stud")
    f_group = fields.Char(string="Группа", compute="_compute_stud")
    docx_file = fields.Binary(string="Справка DOCX", attachment=True)
    docx_file_name = fields.Char(string="Имя файла DOCX", default=f"spravka.docx")
    
    def Method(self):
        for record in self:
            record.display_name = f"справка {record.docx_file_name}"
    
    @api.depends('f_student')
    def _compute_stud(self):
        for record in self:
            record.f_group = record.f_student.f_group.display_name if record.f_student and record.f_student.f_group else False
            record.stud_birth_day = record.f_student.birth_day if record.f_student else False
            record.group_creating_date = record.f_student.group_creating_date if record.f_student else False

    @api.depends('f_student', 'group_creating_date', 'f_group')
    def _compute_docx_file(self):
        for record in self:
            if record.f_student and record.group_creating_date and record.f_group and record.stud_birth_day:
                doc = docx.Document()
                title = doc.add_heading("", level=1)
                title_run = title.add_run("МИНОБРНАУКИ РОССИИ")
                title_run.font.color.rgb = RGBColor(0, 0, 0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                subtitle = doc.add_heading("", level=2)
                subtitle_text = (
                    'Федеральное государственное бюджетное образовательное учреждение высшего образования\n'
                    '«Северо-Восточный Технический Университет имени Петра Великого»\n'
                    '(ФГБОУ ВО «СВТУ им. Петра Великого»)')
                subtitle_run = subtitle.add_run(subtitle_text)
                subtitle_run.font.color.rgb = RGBColor(0, 0, 0)
                subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph('\n\n444444, г. Санкт-Петербург, улица Пушкина, 4').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                spravka_date = fields.Date.today()
                MONTHS = ['января', 'февраля', 'марта', 'апреля', 'мая', 'июня', 'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']
                doc.add_paragraph(f'Дата документа: {spravka_date.day} {MONTHS[int(spravka_date.month)-1]} {spravka_date.year}г.').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
                center_paragraph = doc.add_paragraph()
                center_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = center_paragraph.add_run('\n\n\n\n\nСПРАВКА\n\n')
                run.bold = True

                body = doc.add_paragraph()
                body.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                body.add_run(
                    f'Выдана в том, что {record.f_student.display_name}, '
                    f'{record.stud_birth_day.day}.{record.stud_birth_day.month}.{record.stud_birth_day.year} г.р. '
                    f'является обучающимся группы {record.f_group}. Приказ о зачислении от {record.group_creating_date.day}.{record.group_creating_date.month}.{record.group_creating_date.year} \n')

                body.add_run('\nВыдана для предоставления по месту требования.\n')
                footer = doc.add_paragraph()
                footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                footer.add_run('\n\n\nПроректор по учебной работе\n')
                footer.add_run('И.И. Иванов')
                footer.add_run(f'\n{spravka_date.day} {MONTHS[int(spravka_date.month)-1]} {spravka_date.year}г.\n')
                
                doc_stream = io.BytesIO()
                doc.save(doc_stream)
                doc_stream.seek(0)
                record.docx_file = base64.b64encode(doc_stream.getvalue())
                doc_stream.close()

    def generate_docx(self):
        self._compute_docx_file()
        return {
            'type': 'ir.actions.act_url',
            'url': f"/web/content/?model=new_module.spravka_treb&field=docx_file&id={self.id}&download=true&filename={self.docx_file_name}",
            'target': 'self',
        }